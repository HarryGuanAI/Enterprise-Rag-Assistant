from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def load_document_text(path: str, file_type: str) -> str:
    """把不同格式文档解析为纯文本。

    第一版边界：
    - PDF 仅支持文本型 PDF，不做 OCR。
    - DOCX 支持段落和简单表格文本提取，不支持老 `.doc`。
    - Markdown/TXT 直接读取文本。
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")

    normalized_type = file_type.lower().lstrip(".")
    if normalized_type in {"md", "txt"}:
        return _read_text_file(file_path)

    if normalized_type == "pdf":
        reader = PdfReader(str(file_path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[第 {index} 页]\n{text.strip()}")
        return "\n\n".join(pages)

    if normalized_type == "docx":
        doc = DocxDocument(str(file_path))
        blocks: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)

    raise ValueError(f"不支持的文件类型：{file_type}")

